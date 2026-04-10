from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# Brand constants
FONT_NAME = 'Crimson Pro'
COLOR_BLACK = RGBColor(0, 0, 0)
ACCENT_COLOR_HEX = '003366'

H1_SIZE = Pt(32)
H2_SIZE = Pt(24)
BODY_SIZE = Pt(14)
SMALL_SIZE = Pt(12)

LINE_SPACING = 1.5

SPACE_BEFORE_H1 = Pt(30)
SPACE_AFTER_H1 = Pt(15)
SPACE_BEFORE_H2 = Pt(20)
SPACE_AFTER_H2 = Pt(10)
SPACE_AFTER_BODY = Pt(15)
SPACE_AFTER_LIST = Pt(10)


def classify_paragraph(paragraph):
    """Determine the role of a paragraph: heading1, heading2, list_item, or body."""
    style_name = paragraph.style.name.lower()

    if 'heading 1' in style_name or style_name == 'title':
        return 'heading1'
    if 'heading 2' in style_name or style_name == 'subtitle':
        return 'heading2'
    if 'list' in style_name:
        return 'list_item'

    # Check for bullet/numbered list via XML numbering
    pPr = paragraph._p.find(qn('w:pPr'))
    if pPr is not None and pPr.find(qn('w:numPr')) is not None:
        return 'list_item'

    # Heuristic fallback: detect headings by formatting
    runs_with_text = [r for r in paragraph.runs if r.text.strip()]
    if runs_with_text:
        all_bold = all(r.bold for r in runs_with_text)
        sizes = [r.font.size for r in runs_with_text if r.font.size]
        if all_bold and sizes:
            max_size = max(sizes)
            if max_size >= Pt(24):
                return 'heading1'
            if max_size >= Pt(18):
                return 'heading2'

    return 'body'


def add_accent_line(paragraph):
    """Add a 3pt dark blue bottom border (accent line) under a paragraph."""
    pPr = paragraph._p.get_or_add_pPr()

    # Remove any existing border
    existing = pPr.find(qn('w:pBdr'))
    if existing is not None:
        pPr.remove(existing)

    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '24')  # 24 eighth-points = 3pt
    bottom.set(qn('w:space'), '4')
    bottom.set(qn('w:color'), ACCENT_COLOR_HEX)
    pBdr.append(bottom)
    pPr.append(pBdr)


def remove_accent_line(paragraph):
    """Remove any bottom border from a paragraph."""
    pPr = paragraph._p.find(qn('w:pPr'))
    if pPr is not None:
        existing = pPr.find(qn('w:pBdr'))
        if existing is not None:
            pPr.remove(existing)


def apply_font_to_run(run, size, bold=None):
    """Apply brand font settings to a single run."""
    run.font.name = FONT_NAME
    run.font.size = size
    run.font.color.rgb = COLOR_BLACK
    run.font.italic = False
    if bold is not None:
        run.font.bold = bold

    # Set the font for East Asian and complex script text as well
    rPr = run._r.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.insert(0, rFonts)
    rFonts.set(qn('w:ascii'), FONT_NAME)
    rFonts.set(qn('w:hAnsi'), FONT_NAME)
    rFonts.set(qn('w:cs'), FONT_NAME)


def apply_title_style(paragraph):
    """Style for the title page heading: centered, large, no accent line."""
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pf = paragraph.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after = Pt(10)
    pf.line_spacing = LINE_SPACING
    remove_accent_line(paragraph)
    for run in paragraph.runs:
        apply_font_to_run(run, H1_SIZE, bold=True)


def apply_subtitle_style(paragraph):
    """Style for the title page subtitle: centered, H2 size."""
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pf = paragraph.paragraph_format
    pf.space_before = Pt(5)
    pf.space_after = Pt(10)
    pf.line_spacing = LINE_SPACING
    remove_accent_line(paragraph)
    for run in paragraph.runs:
        apply_font_to_run(run, H2_SIZE, bold=True)


def apply_h1_style(paragraph):
    """Style for content H1 headings: left-aligned with accent line."""
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    pf = paragraph.paragraph_format
    pf.space_before = SPACE_BEFORE_H1
    pf.space_after = SPACE_AFTER_H1
    pf.line_spacing = LINE_SPACING
    add_accent_line(paragraph)
    for run in paragraph.runs:
        apply_font_to_run(run, H1_SIZE, bold=True)


def apply_h2_style(paragraph):
    """Style for H2 subheadings: left-aligned, no accent line."""
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    pf = paragraph.paragraph_format
    pf.space_before = SPACE_BEFORE_H2
    pf.space_after = SPACE_AFTER_H2
    pf.line_spacing = LINE_SPACING
    remove_accent_line(paragraph)
    for run in paragraph.runs:
        apply_font_to_run(run, H2_SIZE, bold=True)


def apply_body_style(paragraph):
    """Style for regular body text."""
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    pf = paragraph.paragraph_format
    pf.space_after = SPACE_AFTER_BODY
    pf.line_spacing = LINE_SPACING
    for run in paragraph.runs:
        apply_font_to_run(run, BODY_SIZE)


def apply_list_style(paragraph):
    """Style for list items: body font with tighter spacing."""
    pf = paragraph.paragraph_format
    pf.space_after = SPACE_AFTER_LIST
    pf.line_spacing = LINE_SPACING
    for run in paragraph.runs:
        apply_font_to_run(run, BODY_SIZE)


def format_table_cells(table):
    """Apply brand fonts to all text in a table."""
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    apply_font_to_run(run, BODY_SIZE)


def process_document(input_path, output_path):
    """Open a .docx file, apply brand formatting, and save to output_path."""
    doc = Document(input_path)

    # Set margins on all sections
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # Classify all paragraphs first
    paragraphs = list(doc.paragraphs)
    first_h1_found = False
    title_done = False

    for i, paragraph in enumerate(paragraphs):
        # Skip empty paragraphs (preserve them but don't classify)
        if not paragraph.text.strip():
            # Still apply font to any runs (e.g., whitespace runs)
            for run in paragraph.runs:
                apply_font_to_run(run, BODY_SIZE)
            paragraph.paragraph_format.line_spacing = LINE_SPACING
            continue

        role = classify_paragraph(paragraph)

        if role == 'heading1' and not first_h1_found:
            # First H1 is the title
            first_h1_found = True
            apply_title_style(paragraph)
        elif role == 'heading1':
            # Subsequent H1s are content headings
            title_done = True
            apply_h1_style(paragraph)
        elif role == 'heading2' and not title_done:
            # H2 before content starts is a subtitle
            apply_subtitle_style(paragraph)
        elif role == 'heading2':
            apply_h2_style(paragraph)
        elif role == 'list_item':
            title_done = True
            apply_list_style(paragraph)
        else:
            title_done = True
            apply_body_style(paragraph)

    # Process tables
    for table in doc.tables:
        format_table_cells(table)

    doc.save(output_path)
